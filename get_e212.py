import requests
from io import BytesIO, StringIO
import docx
from pycountry import countries
from collections import defaultdict
from pprint import pformat

itu_report = ('https://www.itu.int/dms_pub/itu-t/opb/sp/'
              'T-SP-E.212B-2015-MSW-E.docx')
out_file = 'phone_iso3166/e212.py'
names_out_file = 'phone_iso3166/e212_names.py'

r = requests.get(itu_report)
r.raise_for_status()
docxdata = BytesIO(r.content)
doc = docx.Document(docxdata)
# doc = docx.Document('T-SP-E.212B-2014-MSW-E.docx')

# Mapping between ITU names and pycountry names
transform = {
    'Bolivia (Plurinational State of)': 'Bolivia, Plurinational State of',
    'British Virgin Islands': 'Virgin Islands, British',
    'Cabo Verde': 'Cape Verde',
    'Central African Rep.': 'Central African Republic',
    'Czech Rep.': 'Czech Republic',
    'Dem. Rep. of the Congo': 'Congo, The Democratic Republic of the',
    'Dominican Rep.': 'Dominican Republic',
    'French Departments and Territories in the Indian Ocean': 'Réunion',
    'Hong Kong, China': 'Hong Kong',
    'Iran (Islamic Republic of)': 'Iran, Islamic Republic of',
    'Korea (Rep. of)': 'Korea, Republic of',
    'Lao P.D.R.': "Lao People's Democratic Republic",
    'Macao, China': 'Macao',
    'Micronesia': 'Micronesia, Federated States of',
    'Moldova (Republic of)': 'Moldova, Republic of',
    'Tanzania': 'Tanzania, United Republic of',
    'The Former Yugoslav Republic of Macedonia': 'Macedonia, Republic of',
    'Venezuela (Bolivarian Republic of)': 'Venezuela, Bolivarian Republic of',

}

networkdict = defaultdict(lambda: defaultdict(tuple))

# Generate the networkdict from the docx data
# Takes about 20 minutes on a Macbook Pro so be patient
table = doc.tables[1]
for i, row in enumerate(table.rows[1:]):
    # print(i)
    # if i < 0:
    #    continue
    countryname = row.cells[0].paragraphs[0].text
    if countryname in transform:
        countryname = transform[countryname]
    network = row.cells[1].paragraphs[0].text
    mccmnc = row.cells[2].paragraphs[0].text
    if mccmnc:
        mcc, mnc = mccmnc.split(' ')
        try:
            mcc = int(mcc)
            mnc = int(mnc)
            country = countries.get(name=countryname)
            print('{} ({}): {} - {}'.format(countryname, country.alpha2,
                                            mccmnc, network))
            networkdict[mcc][mnc] = (country.alpha2, network)
        except ValueError:
            pass


operators = {k: dict(v) for k, v in networkdict.items()}

countries = defaultdict(list)
for mcc, mncs in operators.items():
    for mnc, c in mncs.items():
        countries[c[0]].append((mcc, mnc))
countries = {k: v for k, v in countries.items()}

oprmap = StringIO()
oprmap.write('#!/usr/bin/env python\n')
oprmap.write('# -*- coding: utf-8 -*-\n\n')
oprmap.write('# Generated by get_e212.py\n')
oprmap.write('# Based on https://www.itu.int/pub/T-SP-E.212B-2015\n\n')
oprmap.write('operators = \\\n')
oprmap.write(pformat(operators))
oprmap.write('\n\n\ncountries = \\\n')
oprmap.write(pformat(countries))

with open(names_out_file, 'w') as f:
    f.write(oprmap.getvalue() + '\n')
print('Wrote ' + names_out_file)

# Construct a new dict, but compressed to if the MCC all share the same country
# we output the country directly else a dict with MNC codes mapped to country
networks = {}
for mcc in operators:
    mncs = operators[mcc]
    sample = next(iter(mncs.values()))[0]
    if all(c == sample for c, n in mncs.values()):
        networks[mcc] = sample
    else:
        networks[mcc] = {n: t[0] for n, t in mncs.items()}

netmap = StringIO()
netmap.write('# Generated by get_e212.py\n')
netmap.write('# Based on https://www.itu.int/pub/T-SP-E.212B-2014\n\n')
netmap.write('networks = \\\n')
netmap.write(pformat(networks))
with open(out_file, 'w') as f:
    f.write(netmap.getvalue() + '\n')
print('Wrote ' + out_file)
